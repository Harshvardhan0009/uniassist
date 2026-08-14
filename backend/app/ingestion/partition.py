"""
Document Partitioning — Step 3 of the pipeline.

Extracts text and structured tables from PDF, DOCX, TXT, and Markdown files.
Uses `pdfplumber` for table-aware PDF parsing so tables are converted into
clean Markdown tables, preserving row and column relationships.
"""

import logging
from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def partition_pdf(file_path: Path) -> list[Document]:
    """
    Extract text and structured tables from a PDF file.
    Converts tables into clean Markdown format.
    """
    logger.info(f"Partitioning PDF (table-aware): {file_path.name}")
    documents = []

    try:
        import pdfplumber

        with pdfplumber.open(str(file_path)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                tables = page.extract_tables()

                content_parts = [page_text]

                if tables:
                    table_mds = []
                    for table in tables:
                        clean_rows = []
                        for row in table:
                            clean_cells = [
                                cell.replace("\n", " ").strip() if cell else ""
                                for cell in row
                            ]
                            if any(clean_cells) and len([c for c in clean_cells if c]) > 1:
                                clean_rows.append(clean_cells)

                        if clean_rows and len(clean_rows) >= 2:
                            header = clean_rows[0]
                            md_table = "| " + " | ".join(header) + " |\n"
                            md_table += "| " + " | ".join(["---"] * len(header)) + " |\n"
                            for r in clean_rows[1:]:
                                r = r + [""] * (len(header) - len(r))
                                md_table += "| " + " | ".join(r[: len(header)]) + " |\n"
                            table_mds.append(md_table)

                    if table_mds:
                        content_parts.append("\n\n### Extracted Table Data:\n" + "\n\n".join(table_mds))

                full_content = "\n\n".join(part for part in content_parts if part.strip())
                if full_content.strip():
                    doc = Document(
                        page_content=full_content,
                        metadata={
                            "source": str(file_path),
                            "filename": file_path.name,
                            "page_number": page_idx + 1,
                            "category": "TableAndText" if tables else "NarrativeText",
                        },
                    )
                    documents.append(doc)

        logger.info(f"  → {len(documents)} pages (with table preservation) from {file_path.name}")
        return documents

    except Exception as e:
        logger.warning(f"pdfplumber extraction failed for {file_path.name}: {e}. Falling back to PyPDFLoader.")
        loader = PyPDFLoader(str(file_path))
        pages = loader.load()
        for page in pages:
            page.metadata["filename"] = file_path.name
            page.metadata["page_number"] = page.metadata.get("page", 0) + 1
            page.metadata["category"] = "NarrativeText"
        return pages


def partition_docx(file_path: Path) -> list[Document]:
    """Extract text and tables from a DOCX file."""
    logger.info(f"Partitioning DOCX: {file_path.name}")
    try:
        import docx

        doc = docx.Document(str(file_path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract table data
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip())
                if row_text:
                    table_lines.append(row_text)
            if table_lines:
                full_text.append("\n".join(table_lines))

        content = "\n\n".join(full_text)
        if not content.strip():
            return []

        doc_obj = Document(
            page_content=content,
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "page_number": 1,
                "category": "NarrativeText",
            },
        )
        return [doc_obj]
    except Exception as e:
        logger.error(f"Failed to read DOCX {file_path.name}: {e}")
        return []


def partition_text(file_path: Path) -> list[Document]:
    """Extract text from a TXT or Markdown file."""
    logger.info(f"Partitioning Text/MD: {file_path.name}")
    try:
        loader = TextLoader(str(file_path), encoding="utf-8")
        docs = loader.load()
        for d in docs:
            d.metadata["filename"] = file_path.name
            d.metadata["page_number"] = 1
            d.metadata["category"] = "NarrativeText"
        return docs
    except Exception as e:
        logger.error(f"Failed to read Text {file_path.name}: {e}")
        return []


def partition_file(file_path: Path) -> list[Document]:
    """Partition a single document based on extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return partition_pdf(file_path)
    elif suffix == ".docx":
        return partition_docx(file_path)
    elif suffix in (".txt", ".md"):
        return partition_text(file_path)
    else:
        return []


def partition_directory(data_dir: Path) -> dict[str, list[Document]]:
    """Extract text from all supported files in a directory (recursive)."""
    supported_extensions = {".pdf", ".docx", ".txt", ".md"}
    all_files = [
        f for f in sorted(data_dir.rglob("*"))
        if f.is_file() and f.suffix.lower() in supported_extensions and not f.name.startswith(".")
    ]

    if not all_files:
        logger.warning(f"No supported document files found in {data_dir}")
        return {}

    logger.info(f"Found {len(all_files)} document files in {data_dir}")

    results: dict[str, list[Document]] = {}
    for file_path in all_files:
        try:
            docs = partition_file(file_path)
            if docs:
                results[file_path.name] = docs
        except Exception as e:
            logger.error(f"Failed to partition {file_path.name}: {e}")

    total = sum(len(v) for v in results.values())
    logger.info(f"Partitioning complete: {total} total document pages from {len(results)} files")
    return results
